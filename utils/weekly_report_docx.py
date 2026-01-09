"""
📄 ULTRA PROFESSIONAL WEEKLY REPORT DOCX GENERATOR
Haftalik ish rejasini Word formatida yaratadi - PROFESSIONAL DIZAYN

Faylni utils/ papkasiga joylashtiring

pip install python-docx --break-system-packages
"""

import logging
from docx import Document
from docx.shared import Inches, Pt, Cm, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

logger = logging.getLogger(__name__)


class WeeklyReportDocx:
    """Ultra Professional haftalik ish rejasi DOCX yaratuvchi"""

    def __init__(self):
        self.days_uz = {
            'dushanba': 'DUSHANBA',
            'seshanba': 'SESHANBA',
            'chorshanba': 'CHORSHANBA',
            'payshanba': 'PAYSHANBA',
            'juma': 'JUMA',
            'shanba': 'SHANBA'
        }

        self.day_themes = {
            'dushanba': 'Yoshlar muammolarini o\'rganish',
            'seshanba': 'Kitobxonlik',
            'chorshanba': 'Tadbirkorlik va bandlik masalalari kuni',
            'payshanba': 'Profilaktika tadbirlari',
            'juma': 'Harbiy vatanparvarlik va Zakovat',
            'shanba': 'Yetakchining shaxsiy tashabbusi'
        }

        # ═══════════════════════════════════════════════════════════
        # 🎨 PROFESSIONAL RANGLAR PALITRASI
        # ═══════════════════════════════════════════════════════════
        self.colors = {
            # Sarlavha qatori - quyuq ko'k
            'header_bg': '1F4E79',
            'header_text': 'FFFFFF',

            # Kun sarlavhalari - o'rtacha ko'k
            'day_header_bg': '2E75B6',
            'day_header_text': 'FFFFFF',

            # Juft qatorlar - och kulrang
            'row_even_bg': 'F2F2F2',

            # Toq qatorlar - oq
            'row_odd_bg': 'FFFFFF',

            # T/r ustuni - maxsus
            'number_bg': 'DEEAF6',

            # Chegaralar
            'border_dark': '1F4E79',
            'border_light': 'BDD7EE',
        }

    def set_cell_border(self, cell, color="000000", size="4", style="single"):
        """Professional yacheyka chegaralari"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Eski borderlarni olib tashlash
        for border in tcPr.findall(qn('w:tcBorders')):
            tcPr.remove(border)

        tcBorders = OxmlElement('w:tcBorders')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), style)
            border.set(qn('w:sz'), size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), color)
            tcBorders.append(border)

        tcPr.append(tcBorders)

    def set_cell_shading(self, cell, color):
        """Yacheyka fon rangini o'rnatish"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Eski shading ni olib tashlash
        for shading in tcPr.findall(qn('w:shd')):
            tcPr.remove(shading)

        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), color)
        tcPr.append(shading)

    def set_cell_margins(self, cell, top=80, bottom=80, left=120, right=120):
        """Yacheyka ichki marginlarini o'rnatish"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Eski marginlarni olib tashlash
        for mar in tcPr.findall(qn('w:tcMar')):
            tcPr.remove(mar)

        tcMar = OxmlElement('w:tcMar')

        for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            margin = OxmlElement(f'w:{margin_name}')
            margin.set(qn('w:w'), str(margin_value))
            margin.set(qn('w:type'), 'dxa')
            tcMar.append(margin)

        tcPr.append(tcMar)

    def set_cell_width(self, cell, width_cm):
        """Yacheyka kengligini aniq o'rnatish"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Eski width ni olib tashlash
        for w in tcPr.findall(qn('w:tcW')):
            tcPr.remove(w)

        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width_cm * 567)))
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

    def set_row_height(self, row, height_cm):
        """Qator balandligini o'rnatish"""
        tr = row._tr
        trPr = tr.get_or_add_trPr()

        # Eski height ni olib tashlash
        for h in trPr.findall(qn('w:trHeight')):
            trPr.remove(h)

        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), str(int(height_cm * 567)))
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    def format_cell_text(self, cell, text, bold=False, size=10, align='left',
                         font='Times New Roman', color=None):
        """Professional matn formatlash"""
        # Barcha paragraphlarni tozalash
        for p in cell.paragraphs:
            p.clear()

        paragraph = cell.paragraphs[0]

        # Alignment
        if align == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Paragraph spacing
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.15

        run = paragraph.add_run(str(text))
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold

        if color:
            run.font.color.rgb = RGBColor.from_string(color)

        # Times New Roman uchun
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font)

        return run

    def add_multi_line_text(self, cell, lines, size=10, font='Times New Roman',
                            align='center', bold_first=False):
        """Ko'p qatorli matn - professional"""
        # Tozalash
        for p in cell.paragraphs:
            p.clear()

        for i, line in enumerate(lines):
            if i == 0:
                paragraph = cell.paragraphs[0]
            else:
                paragraph = cell.add_paragraph()

            # Alignment
            if align == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif align == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.1

            run = paragraph.add_run(str(line))
            run.font.name = font
            run.font.size = Pt(size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font)

            if bold_first and i == 0:
                run.font.bold = True

    def style_header_cell(self, cell, text, width):
        """Sarlavha yacheykasini professional stilizatsiya qilish"""
        self.set_cell_width(cell, width)
        self.set_cell_border(cell, self.colors['border_dark'], "8")
        self.set_cell_shading(cell, self.colors['header_bg'])
        self.set_cell_margins(cell, 100, 100, 120, 120)

        run = self.format_cell_text(cell, text, bold=True, size=11, align='center')
        run.font.color.rgb = RGBColor.from_string(self.colors['header_text'])

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def style_day_header_cell(self, cell, day_name, day_theme, total_width):
        """Kun sarlavhasi - professional stil"""
        self.set_cell_width(cell, total_width)
        self.set_cell_border(cell, self.colors['border_dark'], "6")
        self.set_cell_shading(cell, self.colors['day_header_bg'])
        self.set_cell_margins(cell, 80, 80, 150, 150)

        # Kun nomi
        cell.paragraphs[0].clear()
        para1 = cell.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = para1.add_run(day_name)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor.from_string(self.colors['day_header_text'])
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

        # Kun mavzusi
        para2 = cell.add_paragraph()
        para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = para2.add_run(day_theme)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor.from_string(self.colors['day_header_text'])
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    def style_data_cell(self, cell, width, is_even_row=False, is_number=False):
        """Ma'lumot yacheykasini stilizatsiya qilish"""
        self.set_cell_width(cell, width)
        self.set_cell_border(cell, self.colors['border_light'], "4")
        self.set_cell_margins(cell, 60, 60, 100, 100)

        # Fon rangi
        if is_number:
            self.set_cell_shading(cell, self.colors['number_bg'])
        elif is_even_row:
            self.set_cell_shading(cell, self.colors['row_even_bg'])
        else:
            self.set_cell_shading(cell, self.colors['row_odd_bg'])

        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def create_weekly_report(
            self,
            content: dict,
            output_path: str,
            full_name: str,
            mahalla: str,
            tuman: str,
            week_date: str
    ) -> bool:
        """
        ULTRA PROFESSIONAL haftalik ish rejasi DOCX yaratish
        """

        try:
            doc = Document()

            # ═══════════════════════════════════════════════════════════
            # 📄 SAHIFA SOZLAMALARI - A4 Landscape
            # ═══════════════════════════════════════════════════════════
            section = doc.sections[0]
            section.orientation = 1  # Landscape
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)

            # ═══════════════════════════════════════════════════════════
            # 📝 SARLAVHA - Professional
            # ═══════════════════════════════════════════════════════════

            # Asosiy sarlavha
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_after = Pt(0)

            title_run = title_para.add_run(
                f"MAHALLADAGI YOSHLAR YETAKCHISI {full_name.upper()}NING"
            )
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor.from_string(self.colors['header_bg'])
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Ikkinchi qator
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.paragraph_format.space_before = Pt(0)
            subtitle_para.paragraph_format.space_after = Pt(6)

            subtitle_run = subtitle_para.add_run(
                f"{week_date} DAVRIDAGI HAFTALIK ISH REJASI"
            )
            subtitle_run.font.name = 'Times New Roman'
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.bold = True
            subtitle_run.font.color.rgb = RGBColor.from_string(self.colors['header_bg'])
            subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # Mahalla va tuman
            location_para = doc.add_paragraph()
            location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            location_para.paragraph_format.space_after = Pt(12)

            location_run = location_para.add_run(f"{mahalla}, {tuman}")
            location_run.font.name = 'Times New Roman'
            location_run.font.size = Pt(11)
            location_run.font.italic = True
            location_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # ═══════════════════════════════════════════════════════════
            # 📊 JADVAL - 4 ustun, professional dizayn
            # ═══════════════════════════════════════════════════════════

            # Ustun kengliklari (jami ~27.16 cm)
            col_widths = [1.3, 13.5, 5.5, 6.86]

            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Jadval properties
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

            # 100% kenglik
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            tblPr.append(tblW)

            # Jadval chegarasi
            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:color'), self.colors['border_dark'])
                tblBorders.append(border)
            tblPr.append(tblBorders)

            # ═══════════════════════════════════════════════════════════
            # 🏷️ SARLAVHA QATORI
            # ═══════════════════════════════════════════════════════════
            header_row = table.rows[0]
            headers = ['T/r', 'Loyiha va tadbirlar', 'O\'tkazilish vaqti va joyi', 'Mas\'ul va hamkorlar']

            for i, (header, width) in enumerate(zip(headers, col_widths)):
                self.style_header_cell(header_row.cells[i], header, width)

            self.set_row_height(header_row, 1.0)

            # ═══════════════════════════════════════════════════════════
            # 📅 KUNLAR VA VAZIFALAR
            # ═══════════════════════════════════════════════════════════
            task_number = 1
            row_index = 0

            for day_key, day_name in self.days_uz.items():
                tasks = content.get(day_key, [])

                if not tasks:
                    continue

                # ─────────────────────────────────────────────────────
                # 📌 KUN SARLAVHASI
                # ─────────────────────────────────────────────────────
                day_row = table.add_row()

                # Yacheykalarni birlashtirish
                day_row.cells[0].merge(day_row.cells[3])
                merged_cell = day_row.cells[0]

                day_theme = self.day_themes.get(day_key, '')
                self.style_day_header_cell(merged_cell, day_name, day_theme, sum(col_widths))

                self.set_row_height(day_row, 1.0)

                # ─────────────────────────────────────────────────────
                # 📝 VAZIFALAR
                # ─────────────────────────────────────────────────────
                for task in tasks:
                    row = table.add_row()
                    cells = row.cells
                    is_even = row_index % 2 == 0

                    # T/r - tartib raqami
                    self.style_data_cell(cells[0], col_widths[0], is_even, is_number=True)
                    self.format_cell_text(cells[0], str(task_number), bold=True, size=10, align='center')

                    # Vazifa tavsifi
                    self.style_data_cell(cells[1], col_widths[1], is_even)
                    vazifa = task.get('vazifa', '')
                    self.format_cell_text(cells[1], vazifa, size=10, align='left')

                    # Vaqt va joy
                    self.style_data_cell(cells[2], col_widths[2], is_even)
                    vaqt = task.get('vaqt', '')
                    joy = task.get('joy', 'Mahalla idorasi')
                    hisobot = task.get('hisobot', '')

                    lines = [vaqt, '', joy]
                    if hisobot:
                        lines.append(f'({hisobot})')
                    self.add_multi_line_text(cells[2], lines, size=9, align='center', bold_first=True)

                    # Mas'ullar
                    self.style_data_cell(cells[3], col_widths[3], is_even)
                    masul = task.get('masul', 'Mahalla yoshlar yetakchisi')
                    self.format_cell_text(cells[3], masul, size=10, align='left')

                    task_number += 1
                    row_index += 1

            # ═══════════════════════════════════════════════════════════
            # 📝 IZOH
            # ═══════════════════════════════════════════════════════════
            doc.add_paragraph()

            note_para = doc.add_paragraph()
            note_para.paragraph_format.space_before = Pt(8)

            note_run = note_para.add_run(
                "Izoh: mahalladagi yoshlarning qiziqishlari va mahalla infratuzilmasiga qarab, "
                "mazkur tadbirlar rejasining kunlari yoki vaqtlari o'zgarishi mumkin."
            )
            note_run.font.name = 'Times New Roman'
            note_run.font.size = Pt(9)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(100, 100, 100)
            note_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

            # ═══════════════════════════════════════════════════════════
            # 💾 SAQLASH
            # ═══════════════════════════════════════════════════════════
            doc.save(output_path)
            logger.info(f"✅ Ultra Professional DOCX yaratildi: {output_path}")
            return True

        except Exception as e:
            logger.error(f"❌ DOCX yaratishda xato: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return False