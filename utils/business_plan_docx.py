# utils/business_plan_docx.py
# Professional biznes reja DOCX generator

import logging
import re
from datetime import datetime
from typing import Dict, List, Tuple

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

logger = logging.getLogger(__name__)

# Rang sxemasi
COLOR_PRIMARY = RGBColor(0x1A, 0x73, 0xE8)       # Ko'k
COLOR_DARK = RGBColor(0x0B, 0x1F, 0x4B)           # To'q navy
COLOR_ACCENT = RGBColor(0x0D, 0x47, 0xA1)         # To'q ko'k
COLOR_LIGHT_BG = RGBColor(0xF0, 0xF4, 0xFF)       # Och ko'k fon
COLOR_GRAY = RGBColor(0x60, 0x60, 0x60)           # Kulrang
COLOR_LIGHT_GRAY = RGBColor(0xBD, 0xBD, 0xBD)     # Och kulrang
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)          # Oq
COLOR_GOLD = RGBColor(0xC9, 0xA2, 0x27)           # Oltin aksent

# Jadval ranglar (hex, shading uchun)
HEX_HEADER_BG = "1A73E8"      # Ko'k header
HEX_ALT_ROW = "EEF4FD"         # Alternating och ko'k
HEX_COVER_BG = "0B1F4B"        # Muqova to'q fon


class BusinessPlanDocx:
    """Professional biznes reja DOCX yaratish"""

    def create(self, content: Dict, output_path: str) -> bool:
        try:
            doc = Document()
            self._setup_document(doc)

            business_name = content.get("business_name", "Biznes Reja")
            industry = content.get("industry", "")

            self._add_cover_page(doc, business_name, industry, content)

            doc.add_page_break()
            self._add_toc(doc)

            sections = [
                ("1. IJROIYA XULOSASI", content.get("executive_summary", "")),
                ("2. TASHABBUSKOR VA KOMPANIYA TAVSIFI", content.get("company_description", "")),
                ("3. BOZOR TAHLILI", content.get("market_analysis", "")),
                ("4. MAHSULOT VA XIZMATLAR", content.get("product_service_section") or content.get("product_service", "")),
                ("5. MARKETING VA SAVDO STRATEGIYASI", content.get("marketing_strategy", "")),
                ("6. OPERATSION REJA", content.get("operations_plan", "")),
                ("7. MOLIYAVIY PROGNOZ", content.get("financial_projections", "")),
                ("8. BOSHQARUV JAMOASI VA KADRLAR", content.get("team_section", "")),
                ("9. RISK TAHLILI VA BOSHQARUV", content.get("risk_analysis", "")),
                ("10. XULOSA VA INVESTOR/BANKGA MUROJAAT", content.get("conclusion", "")),
            ]

            for section_title, section_text in sections:
                doc.add_page_break()
                self._add_section(doc, section_title, section_text)

            doc.save(output_path)
            logger.info(f"✅ Biznes reja saqlandi: {output_path}")
            return True

        except Exception as e:
            logger.error(f"❌ Biznes reja DOCX xato: {e}", exc_info=True)
            return False

    def _setup_document(self, doc: Document):
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = Pt(18)

    # ==========================================================================
    # COVER PAGE — qayta dizayn qilingan
    # ==========================================================================

    def _add_cover_page(self, doc: Document, business_name: str, industry: str, content: Dict):
        """Professional muqova — rang fon, katta tipografiya, chegaralar"""

        # Yuqori — sana va "CONFIDENTIAL" brend chizig'i
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Sana: {datetime.now().strftime('%d.%m.%Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_GRAY
        run.font.name = 'Times New Roman'

        # Yuqori dekorativ chiziq (qalin, ko'k)
        self._add_thick_line(doc, COLOR_PRIMARY, thickness=3)

        # Bo'sh joy
        for _ in range(3):
            doc.add_paragraph()

        # Kichik "DOCUMENT TYPE" yorlig'i
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("MAXSUS TAYYORLANGAN HUJJAT")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_GOLD
        run.font.bold = True
        run.font.name = 'Times New Roman'
        # Harf orasidagi bo'shliq (letter spacing)
        self._set_character_spacing(run, 40)

        doc.add_paragraph()

        # Asosiy sarlavha — "BIZNES REJA" — KATTA
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(10)
        p.paragraph_format.line_spacing = 1.15  # Normal stildagi Pt(18) ni bekor qilish
        run = p.add_run("BIZNES REJA")
        run.font.bold = True
        run.font.size = Pt(48)
        run.font.color.rgb = COLOR_DARK
        run.font.name = 'Georgia'
        self._set_character_spacing(run, 80)

        # "BIZNES REJA" ostida yupqa oltin chiziq
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("━━━━━━━━━━━━━━━━━━")
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_GOLD

        doc.add_paragraph()

        # Biznes nomi — rang fonli table cell ichida (chiroyli ko'rinish)
        self._add_title_box(doc, business_name.upper())

        doc.add_paragraph()

        # Soha
        if industry:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(industry)
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_ACCENT
            run.font.italic = True
            run.font.name = 'Times New Roman'

        doc.add_paragraph()
        doc.add_paragraph()

        # Pasport ma'lumotlari — chiroyli info jadval (2 ustun, borderssiz)
        location = content.get("location") or content.get("target_market", "")
        financing = content.get("financing") or content.get("investment", "")
        initiator_type = content.get("initiator_type", "")
        company_info = content.get("company_info", "")
        personal_info = content.get("personal_info", "")

        info_rows = []
        if initiator_type:
            info_rows.append(("Tashabbuskor turi", initiator_type))
        if company_info:
            info_rows.append(("Korxona", company_info))
        if personal_info:
            # Telefon raqamni yashirish
            masked = self._mask_phone(personal_info)
            info_rows.append(("Mas'ul shaxs", masked))
        if location:
            info_rows.append(("Hudud", location))
        if financing:
            info_rows.append(("Moliyalashtirish", financing))

        if info_rows:
            self._add_cover_info_table(doc, info_rows)

        # Pastki qism
        for _ in range(2):
            doc.add_paragraph()

        self._add_thick_line(doc, COLOR_PRIMARY, thickness=3)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("MAXFIY HUJJAT  •  Faqat ishbilarmon maqsadlarda foydalanish uchun")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_GRAY
        run.font.italic = True
        run.font.name = 'Times New Roman'
        self._set_character_spacing(run, 20)

    def _add_title_box(self, doc: Document, title: str):
        """Biznes nomi uchun rangli fon bilan chiroyli box"""
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Cm(14)

        # Fon rangini berish
        self._set_cell_shading(cell, HEX_COVER_BG)
        # Vertikal markazlashtirish
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Matn
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(14)
        p.paragraph_format.line_spacing = 1.2

        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = COLOR_WHITE
        run.font.name = 'Georgia'
        self._set_character_spacing(run, 60)

        # Cell chegaralari — oltin rang
        self._set_cell_borders(cell, color="C9A227", size=8)

    def _add_cover_info_table(self, doc: Document, rows: List[Tuple[str, str]]):
        """Muqovadagi pasport ma'lumotlari — borderssiz ikki ustunli"""
        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        for i, (label, value) in enumerate(rows):
            left = table.cell(i, 0)
            right = table.cell(i, 1)

            left.width = Cm(5)
            right.width = Cm(9)

            # Alternating yumshoq fon
            if i % 2 == 0:
                self._set_cell_shading(left, "F5F8FC")
                self._set_cell_shading(right, "F5F8FC")

            # Label
            p = left.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(f"{label}:")
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_GRAY
            run.font.bold = True
            run.font.name = 'Times New Roman'

            # Value
            p = right.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run(value)
            run.font.size = Pt(11)
            run.font.color.rgb = COLOR_DARK
            run.font.name = 'Times New Roman'

            # Chegarasiz
            self._set_cell_borders(left, color="FFFFFF", size=0)
            self._set_cell_borders(right, color="FFFFFF", size=0)

    # ==========================================================================
    # MUNDARIJA
    # ==========================================================================

    def _add_toc(self, doc: Document):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("MUNDARIJA")
        run.font.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_DARK
        run.font.name = 'Georgia'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_horizontal_line(doc, COLOR_PRIMARY)
        doc.add_paragraph()

        toc_items = [
            ("1. Ijroiya Xulosasi", "3"),
            ("2. Tashabbuskor va Kompaniya Tavsifi", "4"),
            ("3. Bozor Tahlili", "6"),
            ("4. Mahsulot va Xizmatlar", "8"),
            ("5. Marketing va Savdo Strategiyasi", "10"),
            ("6. Operatsion Reja", "12"),
            ("7. Moliyaviy Prognoz", "14"),
            ("8. Boshqaruv Jamoasi va Kadrlar", "17"),
            ("9. Risk Tahlili va Boshqaruv", "19"),
            ("10. Xulosa va Investor/Bankga Murojaat", "21"),
        ]

        for title, page in toc_items:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(f"  {title}")
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run2 = p.add_run(f"{'.' * max(1, 60 - len(title))} {page}")
            run2.font.size = Pt(12)
            run2.font.color.rgb = COLOR_GRAY
            run2.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(4)

    # ==========================================================================
    # BO'LIM — matn va haqiqiy jadvallar
    # ==========================================================================

    def _add_section(self, doc: Document, title: str, text: str):
        # Sarlavha
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_DARK
        run.font.name = 'Georgia'

        self._add_horizontal_line(doc, COLOR_PRIMARY)
        doc.add_paragraph()

        if not text:
            return

        # Matnni bloklarga ajratish (paragraph yoki table)
        blocks = self._parse_blocks(text)

        for block in blocks:
            if block["type"] == "table":
                self._add_real_table(doc, block["rows"])
                doc.add_paragraph()
            else:
                self._add_paragraph(doc, block["text"])

        doc.add_paragraph()

    def _parse_blocks(self, text: str) -> List[Dict]:
        """
        Matnni paragraph va table bloklariga ajratish.
        Markdown jadvali:
            | Ust1 | Ust2 |
            |------|------|
            | A    | B    |
        Separator qator (|---|---|) ixtiyoriy — ketma-ket | bilan 2+ qator ham jadval deb olinadi.
        """
        lines = text.split("\n")
        blocks: List[Dict] = []
        buf_para: List[str] = []
        i = 0
        n = len(lines)

        def flush_para():
            if buf_para:
                joined = "\n".join(buf_para).strip()
                if joined:
                    blocks.append({"type": "paragraph", "text": joined})
                buf_para.clear()

        while i < n:
            line = lines[i]
            stripped = line.strip()

            if self._is_table_line(stripped):
                # Ketma-ket table qatorlarini yig'ish
                table_lines = []
                while i < n and (self._is_table_line(lines[i].strip()) or self._is_separator_line(lines[i].strip())):
                    table_lines.append(lines[i].strip())
                    i += 1

                # Kamida 1 qator bo'lsa jadval deb hisoblanadi (cheksiz sonli ustun)
                rows = self._parse_table_rows(table_lines)
                if rows and len(rows[0]) >= 2:
                    flush_para()
                    blocks.append({"type": "table", "rows": rows})
                else:
                    # Aslida jadval emas — matn sifatida
                    buf_para.extend(table_lines)
                continue

            buf_para.append(line)
            i += 1

        flush_para()
        return blocks

    @staticmethod
    def _is_table_line(line: str) -> bool:
        """| bilan boshlanib va tugaydi, kamida 2 ta | ajratuvchi"""
        if not line:
            return False
        if not line.startswith("|") or not line.endswith("|"):
            return False
        return line.count("|") >= 2

    @staticmethod
    def _is_separator_line(line: str) -> bool:
        """|---|---| ko'rinishidagi ajratuvchi qator"""
        if not line.startswith("|"):
            return False
        inner = line.strip("|").strip()
        cells = [c.strip() for c in inner.split("|")]
        if not cells:
            return False
        return all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c)

    @staticmethod
    def _parse_table_rows(table_lines: List[str]) -> List[List[str]]:
        """Markdown jadval qatorlarini hujayralarga ajratish, separatorni olib tashlab"""
        rows = []
        for line in table_lines:
            if BusinessPlanDocx._is_separator_line(line):
                continue
            inner = line.strip("|")
            cells = [c.strip() for c in inner.split("|")]
            rows.append(cells)

        if not rows:
            return []

        # Barcha qatorlarni bir xil ustun soniga keltirish
        max_cols = max(len(r) for r in rows)
        return [r + [""] * (max_cols - len(r)) for r in rows]

    def _add_real_table(self, doc: Document, rows: List[List[str]]):
        """Haqiqiy Word jadvali — header bold+ko'k fon, alternating, chegaralari"""
        if not rows:
            return

        n_rows = len(rows)
        n_cols = len(rows[0])

        table = doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                cell = table.cell(r_idx, c_idx)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Eski paragraphni tozalab, matn qo'yish (markdown bold **x** yoritish bilan)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)

                clean_text = cell_text.replace("**", "").strip() if r_idx == 0 else cell_text.strip()

                if r_idx == 0:
                    # Header
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(clean_text)
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = COLOR_WHITE
                    run.font.name = 'Times New Roman'
                    self._set_cell_shading(cell, HEX_HEADER_BG)
                else:
                    # Data row
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    self._add_formatted_text_to_paragraph(p, clean_text, size=10)
                    # Alternating row color
                    if r_idx % 2 == 0:
                        self._set_cell_shading(cell, HEX_ALT_ROW)

                # Chegaralar
                self._set_cell_borders(cell, color="A3B8D4", size=4)

    def _add_paragraph(self, doc: Document, text: str):
        """Oddiy matn paragrafi (subheading, bullet, matn)"""
        para_lines = text.split("\n")
        for para_text in para_lines:
            para_text = para_text.strip()
            if not para_text:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(18)

            is_subheading = (
                (para_text.startswith('**') and para_text.endswith('**'))
                or (len(para_text) < 80 and para_text.endswith(':') and not para_text.startswith(('-', '•', '*')))
            )
            is_bullet = para_text.startswith(('-', '•', '*', '→', '✓', '▶'))

            if is_subheading:
                clean = para_text.strip('*').strip()
                run = p.add_run(clean)
                run.font.bold = True
                run.font.size = Pt(13)
                run.font.color.rgb = COLOR_ACCENT
                run.font.name = 'Georgia'
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(10)
            elif is_bullet:
                # Bulletni Word belgisi bilan almashtirish
                clean = re.sub(r'^[-•*→✓▶]\s*', '', para_text)
                run_b = p.add_run("•  ")
                run_b.font.size = Pt(12)
                run_b.font.color.rgb = COLOR_PRIMARY
                run_b.font.bold = True
                self._add_formatted_text_to_paragraph(p, clean, size=12)
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.first_line_indent = Cm(-0.5)
            else:
                self._add_formatted_text_to_paragraph(p, para_text, size=12)

    def _add_formatted_text_to_paragraph(self, paragraph, text: str, size: int = 12):
        """**bold** qismlarni to'g'ri formatlash"""
        parts = text.split('**')
        for i, part in enumerate(parts):
            if not part:
                continue
            run = paragraph.add_run(part)
            run.font.size = Pt(size)
            run.font.name = 'Times New Roman'
            if i % 2 == 1:
                run.font.bold = True

    # ==========================================================================
    # Utility metodlari (XML darajasida shading, border, spacing)
    # ==========================================================================

    @staticmethod
    def _set_cell_shading(cell, hex_color: str):
        """Cellga fon rangi berish"""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

    @staticmethod
    def _set_cell_borders(cell, color: str = "A3B8D4", size: int = 4):
        """Cell chegaralarini sozlash. size=0 bo'lsa chegarasiz."""
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        border_val = "nil" if size == 0 else "single"
        for edge in ("top", "left", "bottom", "right"):
            b = OxmlElement(f'w:{edge}')
            b.set(qn('w:val'), border_val)
            if size > 0:
                b.set(qn('w:sz'), str(size))
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), color)
            tc_borders.append(b)
        tc_pr.append(tc_borders)

    @staticmethod
    def _set_character_spacing(run, value: int):
        """Harf orasidagi bo'shliqni kattalashtirish (1/20 pt)"""
        r_pr = run._element.get_or_add_rPr()
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:val'), str(value))
        r_pr.append(spacing)

    def _add_thick_line(self, doc: Document, color: RGBColor, thickness: int = 3):
        """Qalin gorizontal chiziq (paragraph bottom border orqali)"""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        hex_color = "{:02X}{:02X}{:02X}".format(color[0], color[1], color[2])
        p_pr = p._p.get_or_add_pPr()
        p_borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(thickness * 4))
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), hex_color)
        p_borders.append(bottom)
        p_pr.append(p_borders)

    def _add_horizontal_line(self, doc: Document, color: RGBColor = None):
        """Yupqa gorizontal chiziq"""
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(2)

        col = color if color else COLOR_GRAY
        hex_color = "{:02X}{:02X}{:02X}".format(col[0], col[1], col[2])
        p_pr = p._p.get_or_add_pPr()
        p_borders = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), hex_color)
        p_borders.append(bottom)
        p_pr.append(p_borders)

    @staticmethod
    def _mask_phone(text: str) -> str:
        """Telefon raqamning oxirgi 4 raqamini * bilan yashirish"""
        return re.sub(r'(\d{4})(?=\D*$)', '****', text)
